import pandas as pd
import query_mod as qry
import os
import docx
import docx2pdf
import Treeview_table_mod as tbl


class excelClass:
    def __init__(self, master=None):
        self.sql_query = qry.dbQueries()
        self.treeview = tbl.TreeviewGUI()

    def save_student(self, filename, filepath, date1, date2):
        data, columns = self.sql_query.sort_student_report_bydate_excel(date1, date2)
        df = pd.DataFrame(list(data), columns=columns)
        writer = pd.ExcelWriter(
            filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
        )
        df.to_excel(
            writer,
            sheet_name="Sheet1",
            startrow=12,
            startcol=4,
            header=True,
            index=False,
        )
        worksheet = writer.sheets["Sheet1"]
        worksheet.insert_image("A1", ".\SeekU\STI College Balagtas Logo medium.png")
        worksheet.insert_image("F1", ".\SeekU\SeekU Logotype micro.png")
        worksheet.insert_image("L1", ".\SeekU\SeekU small.png")
        writer.save()

    def save_personnel(self, filename, filepath, date1, date2):
        data, columns = self.sql_query.sort_personnel_report_bydate_excel(date1, date2)
        df = pd.DataFrame(list(data), columns=columns)
        writer = pd.ExcelWriter(
            filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
        )
        df.to_excel(
            writer,
            sheet_name="Sheet1",
            startrow=12,
            startcol=4,
            header=True,
            index=False,
        )
        worksheet = writer.sheets["Sheet1"]
        worksheet.insert_image("A1", ".\SeekU\STI College Balagtas Logo medium.png")
        worksheet.insert_image("F1", ".\SeekU\SeekU Logotype micro.png")
        worksheet.insert_image("L1", ".\SeekU\SeekU small.png")
        writer.save()

    def save_visitor(self, filename, filepath, date1, date2):
        data, columns = self.sql_query.sort_visitor_report_bydate_excel(date1, date2)
        df = pd.DataFrame(list(data), columns=columns)
        writer = pd.ExcelWriter(
            filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
        )
        df.to_excel(
            writer,
            sheet_name="Sheet1",
            startrow=12,
            startcol=4,
            header=True,
            index=False,
        )
        worksheet = writer.sheets["Sheet1"]
        worksheet.insert_image("A1", ".\SeekU\STI College Balagtas Logo medium.png")
        worksheet.insert_image("F1", ".\SeekU\SeekU Logotype micro.png")
        worksheet.insert_image("L1", ".\SeekU\SeekU small.png")
        writer.save()

    def print_student(self, filename, filepath, date1, date2):
        data, columns = self.sql_query.sort_student_report_bydate_excel(date1, date2)
        df = pd.DataFrame(list(data), columns=columns)
        writer = pd.ExcelWriter(
            filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
        )
        df.to_excel(
            writer,
            sheet_name="Sheet1",
            startrow=12,
            startcol=4,
            header=True,
            index=False,
        )
        worksheet = writer.sheets["Sheet1"]
        worksheet.insert_image("A1", ".\SeekU\STI College Balagtas Logo medium.png")
        worksheet.insert_image("F1", ".\SeekU\SeekU Logotype micro.png")
        worksheet.insert_image("L1", ".\SeekU\SeekU small.png")
        writer.save()
        os.startfile(filepath + "/" + filename + ".xlsx")

    def print_personnel(self, filename, filepath, date1, date2):
        data, columns = self.sql_query.sort_personnel_report_bydate_excel(date1, date2)
        df = pd.DataFrame(list(data), columns=columns)
        writer = pd.ExcelWriter(
            filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
        )
        df.to_excel(
            writer,
            sheet_name="Sheet1",
            startrow=12,
            startcol=4,
            header=True,
            index=False,
        )
        worksheet = writer.sheets["Sheet1"]
        worksheet.insert_image("A1", ".\SeekU\STI College Balagtas Logo medium.png")
        worksheet.insert_image("F1", ".\SeekU\SeekU Logotype micro.png")
        worksheet.insert_image("L1", ".\SeekU\SeekU small.png")
        writer.save()
        os.startfile(filepath + "/" + filename + ".xlsx")

    def print_visitor(self, filename, filepath, date1, date2):
        data, columns = self.sql_query.sort_visitor_report_bydate_excel(date1, date2)
        df = pd.DataFrame(list(data), columns=columns)
        writer = pd.ExcelWriter(
            filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
        )
        df.to_excel(
            writer,
            sheet_name="Sheet1",
            startrow=12,
            startcol=4,
            header=True,
            index=False,
        )
        worksheet = writer.sheets["Sheet1"]
        worksheet.insert_image("A1", ".\SeekU\STI College Balagtas Logo medium.png")
        worksheet.insert_image("F1", ".\SeekU\SeekU Logotype micro.png")
        worksheet.insert_image("L1", ".\SeekU\SeekU small.png")
        writer.save()
        os.startfile(filepath + "/" + filename + ".xlsx")


class docxClass:
    def __init__(self, master=None):
        self.sql_query = qry.dbQueries()
        self.treeview = tbl.TreeviewGUI()

    # JOCRIAUS--------------------------------------------------------------------------
    def print_word_student_doc(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_student_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.student_report_tree.get_children():
            values = self.treeview.student_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
            row_cells[7].text = row[7]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")
        os.startfile(filepath + "/" + filename + ".docx")

    def save_word_student_doc(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_student_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.student_report_tree.get_children():
            values = self.treeview.student_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
            row_cells[7].text = row[7]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

    # JOCRIAUS--------------------------------------------------------------------------

    def print_word_personnel_doc(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_personnel_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.personnel_report_tree.get_children():
            values = self.treeview.personnel_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")
        os.startfile(filepath + "/" + filename + ".docx")

    def save_word_personnel_doc(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_personnel_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.personnel_report_tree.get_children():
            values = self.treeview.personnel_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

    def print_word_visitor_doc(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_visitor_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.visitor_report_tree.get_children():
            values = self.treeview.visitor_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")
        os.startfile(filepath + "/" + filename + ".docx")

    def save_word_visitor_doc(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_visitor_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.visitor_report_tree.get_children():
            values = self.treeview.visitor_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

    def print_pdf_student(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_student_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.student_report_tree.get_children():
            values = self.treeview.student_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
            row_cells[7].text = row[7]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

        docx_file = filepath + "/" + filename + ".docx"
        pdf_file = filepath + "/" + filename + ".pdf"

        docx2pdf.convert(docx_file, pdf_file)
        os.startfile(filepath + "/" + filename + ".pdf")

    def save_pdf_student(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_student_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.student_report_tree.get_children():
            values = self.treeview.student_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
            row_cells[7].text = row[7]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

        docx_file = filepath + "/" + filename + ".docx"
        pdf_file = filepath + "/" + filename + ".pdf"

        docx2pdf.convert(docx_file, pdf_file)

    def print_pdf_personnel(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_personnel_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.personnel_report_tree.get_children():
            values = self.treeview.personnel_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

        docx_file = filepath + "/" + filename + ".docx"
        pdf_file = filepath + "/" + filename + ".pdf"

        docx2pdf.convert(docx_file, pdf_file)
        os.startfile(filepath + "/" + filename + ".pdf")

    def save_pdf_personnel(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_personnel_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.personnel_report_tree.get_children():
            values = self.treeview.personnel_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
            row_cells[6].text = row[6]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

        docx_file = filepath + "/" + filename + ".docx"
        pdf_file = filepath + "/" + filename + ".pdf"

        docx2pdf.convert(docx_file, pdf_file)

    def print_pdf_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_visitor_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.visitor_report_tree.get_children():
            values = self.treeview.visitor_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

        docx_file = filepath + "/" + filename + ".docx"
        pdf_file = filepath + "/" + filename + ".pdf"

        docx2pdf.convert(docx_file, pdf_file)
        os.startfile(filepath + "/" + filename + ".pdf")

    def save_pdf_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        self.treeview.populate_visitor_report_bydate(date1, date2)
        # open an existing Word document
        doc = docx.Document("../Documents/Document_temp/Report Template.docx")

        # get the first paragraph in the document
        p1 = doc.paragraphs[8]

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        treeview_data = []
        # get the table data of student
        for child in self.treeview.visitor_report_tree.get_children():
            values = self.treeview.visitor_report_tree.item(child)["values"]
            treeview_data.append(values)
        # Inserts the table data of student
        for row in treeview_data:
            row_cells = table.add_row().cells
            row_cells[0].text = row[0]
            row_cells[1].text = row[1]
            row_cells[2].text = row[2]
            row_cells[3].text = row[3]
            row_cells[4].text = row[4]
            row_cells[5].text = row[5]
        # saves the doc to a new file path
        doc.save(filepath + "/" + filename + ".docx")

        docx_file = filepath + "/" + filename + ".docx"
        pdf_file = filepath + "/" + filename + ".pdf"

        docx2pdf.convert(docx_file, pdf_file)

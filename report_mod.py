import tkinter as Tk
import pandas as pd
import query_mod as qry
import os
import docx
import docx.shared
import docx2pdf
import Treeview_table_mod as tbl
import tkinter.messagebox as messbx
from datetime import datetime


class excelClass:
    def __init__(self, master=None):
        self.sql_query = qry.dbQueries()
        self.treeview = tbl.TreeviewGUI()
        self.current_date = datetime.now().date()

    def save_student(self, filepath, filename, date1, date2, section, ufname, ulname):
        data, columns = self.sql_query.sort_student_report_bydate_excel(
            date1, date2, section
        )
        columns = [
            "Student No.",
            "First Name",
            "Last Name",
            "Program",
            "Section",
            "Attendance Date",
            "Time in",
            "Time out",
        ]
        
        generateby_text = "Generate by: " + ufname + " " + ulname
        generateon_text = "Generate on: " + self.current_date.strftime("%B %d, %Y")
        report_count = self.sql_query.student_report_counter(date1, date2, section)
        total_report_text = "Total Generated Records: " + str(report_count)

        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")

            worksheet = writer.sheets["Sheet1"]
            worksheet.write("F9", str(starting_date)+" to " + str(end_date))
            worksheet.insert_image("B1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Student Report.png")
            # worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("L1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)
            # Add the information after printing the table
            startrow = 14
            last_row_index = len(df.index) + startrow  # Assuming start_row is defined

            worksheet.write(last_row_index, start_column, generateby_text)
            worksheet.write(last_row_index + 1, start_column, generateon_text)
            worksheet.write(last_row_index + 2, start_column, total_report_text)

            writer.close()
            os.startfile(filepath + "/" + filename + ".xlsx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    def save_personnel(self, filepath, filename, date1, date2, ptype, ufname, ulname):
        data, columns = self.sql_query.sort_personnel_report_bydate_excel(
            date1, date2, ptype
        )
        columns = [
            "Personnel No.",
            "First Name",
            "Last Name",
            "Personnel Type",
            "Attendance Date",
            "Time in",
            "Time out",
        ]

        generateby_text = "Generate by: " + ufname + " " + ulname
        generateon_text = "Generate on: " + self.current_date.strftime("%B %d, %Y")
        report_count = self.sql_query.personnel_report_counter(date1, date2, ptype)
        total_report_text = "Total Generated Records: " + str(report_count)

        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")

            worksheet = writer.sheets["Sheet1"]
            worksheet.write("F9", str(starting_date)+" to " + str(end_date))
            worksheet.insert_image("B1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Personnel Report.png")
            # worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("K1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)

            # Add the information after printing the table
            startrow = 14
            last_row_index = len(df.index) + startrow  # Assuming start_row is defined

            worksheet.write(last_row_index, start_column, generateby_text)
            worksheet.write(last_row_index + 1, start_column, generateon_text)
            worksheet.write(last_row_index + 2, start_column, total_report_text)

            writer.close()
            os.startfile(filepath + "/" + filename + ".xlsx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    def save_visitor(self, filepath, filename, date1, date2, ufname, ulname):
        data, columns = self.sql_query.sort_visitor_report_bydate_excel(date1, date2)
        columns = [
            "Visitor No.",
            "First Name",
            "Last Name",
            "Attendance Date",
            "Time in",
            "Time out",
        ]

        generateby_text = "Generate by: " + ufname + " " + ulname
        generateon_text = "Generate on: " + self.current_date.strftime("%B %d, %Y")
        report_count = self.sql_query.visitor_report_counter(date1, date2)
        total_report_text = "Total Generated Records: " + str(report_count)

        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 4
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")

            worksheet = writer.sheets["Sheet1"]
            worksheet.write("G9", str(starting_date)+" to " + str(end_date))
            worksheet.insert_image("B1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Visitor Report.png")
            # worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("K1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)

            # Add the information after printing the table
            startrow = 14
            last_row_index = len(df.index) + startrow  # Assuming start_row is defined

            worksheet.write(last_row_index, start_column, generateby_text)
            worksheet.write(last_row_index + 1, start_column, generateon_text)
            worksheet.write(last_row_index + 2, start_column, total_report_text)
            writer.close()
            os.startfile(filepath + "/" + filename + ".xlsx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    """
    def print_student(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_student_report_bydate_excel(date1, date2)
        columns = [
            "Student No.",
            "First Name",
            "Last Name",
            "Program",
            "Section",
            "Attendance Date",
            "Time in",
            "Time out",
        ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 2
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]

            worksheet = writer.sheets["Sheet1"]
            worksheet.write("E9", starting_date)
            worksheet.write("G9", end_date)
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Student Report.png")
            worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("I1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)

            writer.close()
            os.startfile(filepath + "/" + filename + ".xlsx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    def print_personnel(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_personnel_report_bydate_excel(date1, date2)
        columns = [
            "Personnel No.",
            "First Name",
            "Last Name",
            "Personnel Type",
            "Attendance Date",
            "Time in",
            "Time out",
        ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]

            worksheet = writer.sheets["Sheet1"]
            worksheet.write("F9", starting_date)
            worksheet.write("G9", end_date)
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("F1", ".\SeekU\Personnel Report.png")
            worksheet.insert_image("F6", ".\SeekU\Dates.png")
            worksheet.insert_image("J1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)
            writer.close()
            os.startfile(filepath + "/" + filename + ".xlsx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    def print_visitor(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_visitor_report_bydate_excel(date1, date2)
        columns = [
            "Visitor No.",
            "First Name",
            "Last Name",
            "Attendance Date",
            "Time in",
            "Time out",
        ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]

            worksheet = writer.sheets["Sheet1"]
            worksheet.write("E9", starting_date)
            worksheet.write("H9", end_date)
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Visitor Report.png")
            worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("I1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)
            writer.close()
            os.startfile(filepath + "/" + filename + ".xlsx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )
    """


class docxClass:
    def __init__(self, master=None):
        self.sql_query = qry.dbQueries()
        self.treeview = tbl.TreeviewGUI()
        self.main = Tk.Toplevel()
        self.treeview.student_report_treeview(self.main)
        self.treeview.visitor_report_treeview(self.main)
        self.treeview.personnel_report_treeview(self.main)
        self.main.withdraw()

    """
    def print_doc_student(self, filepath, filename, date1, date2,section, ufname, ulname):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2,section)
        # open an existing Word document
        if cont_inue:
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Student Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
            text = str(start_date) + " to " + str(end_date)
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(18)

            space = "\n"
            doc.add_paragraph(space)

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Student Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = str(date1)[0:10]
                    end_date = str(date2)[0:10]
                    text = "Date: " + str(start_date) + " to " + str(end_date)
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
    """

    def save_doc_student(
        self, filepath, filename, date1, date2, section, ufname, ulname
    ):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2, section)
        print(cont_inue)
        # open an existing Word document
        if cont_inue:
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Student Attendance Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")
            text = start_date + " to " + end_date
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(18)

            space = "\n"
            doc.add_paragraph(space)

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Student Attendance Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = date1.strftime("%B %d, %Y")
                    end_date = date2.strftime("%B %d, %Y")
                    text = start_date + " to " + end_date
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            space = "\n"
            doc.add_paragraph(space)
            current_date = datetime.now().date()

            generateby_text = "Generate by: " + ufname + " " + ulname
            generateby_add = doc.add_paragraph(generateby_text)
            generateby_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateby_font = generateby_add.runs[0].font
            generateby_font.size = docx.shared.Pt(15)

            generateon_text = "Generate on: " + current_date.strftime("%B %d, %Y")
            generateon_add = doc.add_paragraph(generateon_text)
            generateon_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateon_font = generateon_add.runs[0].font
            generateon_font.size = docx.shared.Pt(15)

            report_count = self.sql_query.student_report_counter(date1, date2, section)
            total_report_text = "Total Generated Records: " + str(report_count)
            total_report_add = doc.add_paragraph(total_report_text)
            total_report_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            total_report_font = total_report_add.runs[0].font
            total_report_font.size = docx.shared.Pt(15)

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )

    # JOCRIAUS--------------------------------------------------------------------------
    """
    def print_doc_personnel(self, filepath, filename, date1, date2,ptype, ufname, ulname):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2, ptype)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Personnel Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
            text = str(start_date) + " to " + str(end_date)
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Personnel Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = str(date1)[0:10]
                    end_date = str(date2)[0:10]
                    text = "Date: " + str(start_date) + " to " + str(end_date)
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
    """

    def save_doc_personnel(
        self, filepath, filename, date1, date2, ptype, ufname, ulname
    ):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2, ptype)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Personnel Attendance Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")
            text = start_date + " to " + end_date
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Personnel Attendance Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = date1.strftime("%B %d, %Y")
                    end_date = date2.strftime("%B %d, %Y")
                    text = start_date + " to " + end_date
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            space = "\n"
            doc.add_paragraph(space)
            current_date = datetime.now().date()

            generateby_text = "Generate by: " + ufname + " " + ulname
            generateby_add = doc.add_paragraph(generateby_text)
            generateby_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateby_font = generateby_add.runs[0].font
            generateby_font.size = docx.shared.Pt(15)

            generateon_text = "Generate on: " + current_date.strftime("%B %d, %Y")
            generateon_add = doc.add_paragraph(generateon_text)
            generateon_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateon_font = generateon_add.runs[0].font
            generateon_font.size = docx.shared.Pt(15)

            report_count = self.sql_query.personnel_report_counter(date1, date2, ptype)
            total_report_text = "Total Generated Records: " + str(report_count)
            total_report_add = doc.add_paragraph(total_report_text)
            total_report_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            total_report_font = total_report_add.runs[0].font
            total_report_font.size = docx.shared.Pt(15)

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )

    """
    def print_doc_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Visitor Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
            text = str(start_date) + " to " + str(end_date)
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Visitor Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = str(date1)[0:10]
                    end_date = str(date2)[0:10]
                    text = "Date: " + str(start_date) + " to " + str(end_date)
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
    """

    def save_doc_visitor(self, filepath, filename, date1, date2, ufname, ulname):
        # populate the report tree
        cont_inue = self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Visitor Attendance Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")
            text = start_date + " to " + end_date
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Visitor Attendance Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = date1.strftime("%B %d, %Y")
                    end_date = date2.strftime("%B %d, %Y")
                    text = start_date + " to " + end_date
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            space = "\n"
            doc.add_paragraph(space)
            current_date = datetime.now().date()

            generateby_text = "Generate by: " + ufname + " " + ulname
            generateby_add = doc.add_paragraph(generateby_text)
            generateby_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateby_font = generateby_add.runs[0].font
            generateby_font.size = docx.shared.Pt(15)

            generateon_text = "Generate on: " + current_date.strftime("%B %d, %Y")
            generateon_add = doc.add_paragraph(generateon_text)
            generateon_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateon_font = generateon_add.runs[0].font
            generateon_font.size = docx.shared.Pt(15)

            report_count = self.sql_query.visitor_report_counter(date1, date2)
            total_report_text = "Total Generated Records: " + str(report_count)
            total_report_add = doc.add_paragraph(total_report_text)
            total_report_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            total_report_font = total_report_add.runs[0].font
            total_report_font.size = docx.shared.Pt(15)

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )

    """
    def print_pdf_student(self, filepath, filename, date1, date2, section):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2, section)
        # open an existing Word document
        if cont_inue:
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Student Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
            text = str(start_date) + " to " + str(end_date)
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(18)

            space = "\n"
            doc.add_paragraph(space)

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Student Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = str(date1)[0:10]
                    end_date = str(date2)[0:10]
                    text = "Date: " + str(start_date) + " to " + str(end_date)
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
    """

    def save_pdf_student(
        self, filepath, filename, date1, date2, section, ufname, ulname
    ):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2, section)
        print(cont_inue)
        # open an existing Word document
        if cont_inue:
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Student Attendance Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")
            text = start_date + " to " + end_date
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(18)

            space = "\n"
            doc.add_paragraph(space)

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Student Attendance Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = date1.strftime("%B %d, %Y")
                    end_date = date2.strftime("%B %d, %Y")
                    text = start_date + " to " + end_date
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            space = "\n"
            doc.add_paragraph(space)
            current_date = datetime.now().date()

            generateby_text = "Generate by: " + ufname + " " + ulname
            generateby_add = doc.add_paragraph(generateby_text)
            generateby_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateby_font = generateby_add.runs[0].font
            generateby_font.size = docx.shared.Pt(15)

            generateon_text = "Generate on: " + current_date.strftime("%B %d, %Y")
            generateon_add = doc.add_paragraph(generateon_text)
            generateon_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateon_font = generateon_add.runs[0].font
            generateon_font.size = docx.shared.Pt(15)

            report_count = self.sql_query.student_report_counter(date1, date2, section)
            total_report_text = "Total Generated Records: " + str(report_count)
            total_report_add = doc.add_paragraph(total_report_text)
            total_report_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            total_report_font = total_report_add.runs[0].font
            total_report_font.size = docx.shared.Pt(15)

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )

        """
    def print_pdf_personnel(self, filepath, filename, date1, date2, ptype, ufname, ulname):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2, ptype)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Personnel Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
            text = str(start_date) + " to " + str(end_date)
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Personnel Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = str(date1)[0:10]
                    end_date = str(date2)[0:10]
                    text = "Date: " + str(start_date) + " to " + str(end_date)
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        """

    def save_pdf_personnel(
        self, filepath, filename, date1, date2, ptype, ufname, ulname
    ):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2, ptype)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Personnel Attendance Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")
            text = start_date + " to " + end_date
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Personnel Attendance Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = date1.strftime("%B %d, %Y")
                    end_date = date2.strftime("%B %d, %Y")
                    text = start_date + " to " + end_date
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            space = "\n"
            doc.add_paragraph(space)
            current_date = datetime.now().date()

            generateby_text = "Generate by: " + ufname + " " + ulname
            generateby_add = doc.add_paragraph(generateby_text)
            generateby_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateby_font = generateby_add.runs[0].font
            generateby_font.size = docx.shared.Pt(15)

            generateon_text = "Generate on: " + current_date.strftime("%B %d, %Y")
            generateon_add = doc.add_paragraph(generateon_text)
            generateon_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateon_font = generateon_add.runs[0].font
            generateon_font.size = docx.shared.Pt(15)

            report_count = self.sql_query.personnel_report_counter(date1, date2, ptype)
            total_report_text = "Total Generated Records: " + str(report_count)
            total_report_add = doc.add_paragraph(total_report_text)
            total_report_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            total_report_font = total_report_add.runs[0].font
            total_report_font.size = docx.shared.Pt(15)

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )

    """
    def print_pdf_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Visitor Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
            text = str(start_date) + " to " + str(end_date)
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Visitor Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = str(date1)[0:10]
                    end_date = str(date2)[0:10]
                    text = "Date: " + str(start_date) + " to " + str(end_date)
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
        """

    def save_pdf_visitor(self, filepath, filename, date1, date2, ufname, ulname):
        # populate the report tree
        cont_inue = self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Report-Template.docx")
            header_text = "Visitor Attendance Report"
            header_add = doc.add_paragraph(header_text)
            header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            header_font = header_add.runs[0].font
            header_font.size = docx.shared.Pt(24)
            start_date = date1.strftime("%B %d, %Y")
            end_date = date2.strftime("%B %d, %Y")
            text = start_date + " to " + end_date
            date_time = doc.add_paragraph(text)
            date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            font = date_time.runs[0].font
            font.size = docx.shared.Pt(15)

            space = "\n"
            doc.add_paragraph(space)
            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            column_count = len(column_names)
            for idx, row in enumerate(treeview_data, start=1):
                if idx % 18 == 0:
                    doc.add_page_break()
                    space = "\n"
                    doc.add_paragraph(space)
                    header_text = "Visitor Attendance Report"
                    header_add = doc.add_paragraph(header_text)
                    header_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    header_font = header_add.runs[0].font
                    header_font.size = docx.shared.Pt(24)
                    start_date = date1.strftime("%B %d, %Y")
                    end_date = date2.strftime("%B %d, %Y")
                    text = start_date + " to " + end_date
                    date_time = doc.add_paragraph(text)
                    date_time.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

                    font = date_time.runs[0].font
                    font.size = docx.shared.Pt(18)
                    space = "\n"
                    doc.add_paragraph(space)
                    table = doc.add_table(rows=1, cols=column_count)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells

                    for i in range(len(column_names)):
                        hdr_cells[i].text = column_names[i]

                row_cells = table.add_row().cells
                for i in range(column_count):
                    row_cells[i].text = str(row[i])

            space = "\n"
            doc.add_paragraph(space)
            current_date = datetime.now().date()

            generateby_text = "Generate by: " + ufname + " " + ulname
            generateby_add = doc.add_paragraph(generateby_text)
            generateby_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateby_font = generateby_add.runs[0].font
            generateby_font.size = docx.shared.Pt(15)

            generateon_text = "Generate on: " + current_date.strftime("%B %d, %Y")
            generateon_add = doc.add_paragraph(generateon_text)
            generateon_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            generateon_font = generateon_add.runs[0].font
            generateon_font.size = docx.shared.Pt(15)

            report_count = self.sql_query.visitor_report_counter(date1, date2)
            total_report_text = "Total Generated Records: " + str(report_count)
            total_report_add = doc.add_paragraph(total_report_text)
            total_report_add.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            total_report_font = total_report_add.runs[0].font
            total_report_font.size = docx.shared.Pt(15)

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")
            messbx.showinfo(
                "Generated", "The reports have been successfully generated."
            )
